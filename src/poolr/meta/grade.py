"""
GRADE (Grading of Recommendations Assessment, Development and Evaluation)
Evidence table generation for systematic reviews
"""

import json
from dataclasses import dataclass
from enum import Enum
from typing import Any, Dict, List


class GRADEDomain(Enum):
    """GRADE domains for certainty assessment"""

    RISK_OF_BIAS = "Risk of Bias"
    INCONSISTENCY = "Inconsistency"
    INDIRECTNESS = "Indirectness"
    IMPRECISION = "Imprecision"
    PUBLICATION_BIAS = "Publication Bias"


class CertaintyLevel(Enum):
    """GRADE certainty levels"""

    HIGH = "High"
    MODERATE = "Moderate"
    LOW = "Low"
    VERY_LOW = "Very Low"


@dataclass
class GRADEAssessment:
    """Assessment for a single outcome"""

    outcome: str
    studies: int = 0
    design: str = "RCT"  # or "Observational"

    # Domain judgments
    risk_of_bias: str = "Not serious"  # "Not serious", "Serious", "Very serious"
    inconsistency: str = "Not serious"
    indirectness: str = "Not serious"
    imprecision: str = "Not serious"
    publication_bias: str = "Not serious"

    # Starting certainty
    # RCTs start at High, Observational start at Low
    @property
    def starting_certainty(self) -> CertaintyLevel:
        if self.design.upper() in ["RCT", "RANDOMIZED"]:
            return CertaintyLevel.HIGH
        return CertaintyLevel.LOW

    @property
    def final_certainty(self) -> CertaintyLevel:
        """Calculate final certainty based on downgrades"""
        downgrades = 0

        for domain in [
            self.risk_of_bias,
            self.inconsistency,
            self.indirectness,
            self.imprecision,
            self.publication_bias,
        ]:
            if domain == "Serious":
                downgrades += 1
            elif domain == "Very serious":
                downgrades += 2

        # Apply downgrades
        level = self.starting_certainty
        levels = [CertaintyLevel.HIGH, CertaintyLevel.MODERATE, CertaintyLevel.LOW, CertaintyLevel.VERY_LOW]
        idx = levels.index(level)
        idx = min(idx + downgrades, len(levels) - 1)

        # Can upgrade (rare) - not implemented here
        return levels[idx]

    @property
    def downgrade_summary(self) -> List[str]:
        """Return list of domains that caused downgrades"""
        reasons = []
        for domain_name, judgment in [
            ("Risk of bias", self.risk_of_bias),
            ("Inconsistency", self.inconsistency),
            ("Indirectness", self.indirectness),
            ("Imprecision", self.imprecision),
            ("Publication bias", self.publication_bias),
        ]:
            if judgment in ["Serious", "Very serious"]:
                reasons.append(f"{domain_name}: {judgment}")
        return reasons


class GRADETableGenerator:
    """Generate GRADE evidence tables"""

    def __init__(self, project_data: Dict[str, Any]):
        self.project_data = project_data
        self.assessments: List[GRADEAssessment] = []

    def add_outcome(self, outcome: str, studies: int = 0, design: str = "RCT"):
        """Add an outcome for GRADE assessment"""
        assessment = GRADEAssessment(outcome=outcome, studies=studies, design=design)
        self.assessments.append(assessment)
        return assessment

    def auto_populate_from_meta(self):
        """Auto-populate from meta-analysis results"""
        meta = self.project_data.get("meta", {}).get("results", {})
        if not meta:
            return

        studies = meta.get("studies", [])
        k = len(studies)

        # Determine design
        designs = set(s.get("design", "") for s in studies)
        design = "RCT" if any("RCT" in d for d in designs) else "Observational"

        # Get outcomes from extraction
        outcomes = set()
        for s in self.project_data.get("extraction", {}).get("studies", []):
            if s.get("primary_outcome"):
                outcomes.add(s["primary_outcome"])
            if s.get("secondary_outcomes"):
                for o in s["secondary_outcomes"].split("; "):
                    outcomes.add(o)

        # If no outcomes specified, use generic
        if not outcomes:
            outcomes = ["Primary outcome", "Secondary outcome"]

        for outcome in outcomes:
            self.add_outcome(outcome, k, design)

        # Auto-assess domains based on meta-analysis results
        for assessment in self.assessments:
            self._auto_assess_domains(assessment, meta)

    def _auto_assess_domains(self, assessment: GRADEAssessment, meta: Dict[str, Any]):
        """Auto-assess GRADE domains from meta-analysis results"""
        hetero = meta.get("heterogeneity") or {}
        pub_bias = meta.get("publication_bias") or {}
        pooled = meta.get("pooled") or {}

        # Risk of bias - check RoB assessments
        rob = self.project_data.get("rob", {}).get("assessments", [])
        high_rob = sum(1 for a in rob if a.get("data", {}).get("overall") == "High")
        some_rob = sum(1 for a in rob if a.get("data", {}).get("overall") == "Some concerns")
        total_rob = len(rob)

        if total_rob > 0:
            if high_rob / total_rob > 0.5:
                assessment.risk_of_bias = "Serious"
            elif (high_rob + some_rob) / total_rob > 0.5:
                assessment.risk_of_bias = "Serious"
            elif high_rob > 0 or some_rob > 0:
                assessment.risk_of_bias = "Not serious"  # but could be serious

        # Inconsistency - based on I²
        i2 = hetero.get("i2", 0)
        if i2 > 75:
            assessment.inconsistency = "Serious"
        elif i2 > 50:
            assessment.inconsistency = "Serious"
        elif i2 > 25:
            assessment.inconsistency = "Not serious"

        # Imprecision - based on CI width and sample size
        ci_width = pooled.get("ci_upper", 0) - pooled.get("ci_lower", 0)
        effect = pooled.get("effect", 1)
        if effect > 0:
            rel_width = ci_width / effect
            if rel_width > 1.0:  # Very wide CI
                assessment.imprecision = "Serious"
            elif rel_width > 0.5:
                assessment.imprecision = "Not serious"

        # Also consider sample size
        total_n = sum(s.get("int_n", 0) + s.get("ctrl_n", 0) for s in meta.get("studies", []))
        if total_n < 100:
            assessment.imprecision = "Serious"
        elif total_n < 400:
            assessment.imprecision = "Not serious"

        # Publication bias
        if pub_bias:
            egger = pub_bias.get("egger", {})
            if egger.get("significant", False):
                assessment.publication_bias = "Serious"

    def to_dict(self) -> List[Dict[str, Any]]:
        """Convert to list of dicts for export"""
        result = []
        for a in self.assessments:
            result.append(
                {
                    "outcome": a.outcome,
                    "studies": a.studies,
                    "design": a.design,
                    "starting_certainty": a.starting_certainty.value,
                    "final_certainty": a.final_certainty.value,
                    "risk_of_bias": a.risk_of_bias,
                    "inconsistency": a.inconsistency,
                    "indirectness": a.indirectness,
                    "imprecision": a.imprecision,
                    "publication_bias": a.publication_bias,
                    "downgrade_reasons": a.downgrade_summary,
                }
            )
        return result

    def generate_word_table(self, doc) -> None:
        """Add GRADE evidence table to Word document"""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import Cm, Pt

        if not self.assessments:
            return

        doc.add_heading("GRADE Evidence Profile", level=2)

        # Table header
        headers = [
            "Outcome",
            "Studies",
            "Design",
            "Risk of Bias",
            "Inconsistency",
            "Indirectness",
            "Imprecision",
            "Publication Bias",
            "Starting Certainty",
            "Final Certainty",
            "Downgrade Reasons",
        ]

        table = doc.add_table(rows=len(self.assessments) + 1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        # Data rows
        certainty_colors = {"High": "00B050", "Moderate": "FFC000", "Low": "FF6600", "Very Low": "FF0000"}

        for row_idx, a in enumerate(self.assessments):
            row = table.rows[row_idx + 1]
            cells = [
                a.outcome,
                str(a.studies),
                a.design,
                a.risk_of_bias,
                a.inconsistency,
                a.indirectness,
                a.imprecision,
                a.publication_bias,
                a.starting_certainty.value,
                a.final_certainty.value,
                "; ".join(a.downgrade_summary) if a.downgrade_summary else "None",
            ]

            for col_idx, value in enumerate(cells):
                cell = row.cells[col_idx]
                cell.text = str(value)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)

                # Color final certainty
                if col_idx == 9:  # Final certainty column
                    color = certainty_colors.get(value, "000000")
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn

                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), color)
                    cell._tc.get_or_add_tcPr().append(shading)

        # Set column widths
        widths = [Cm(3), Cm(1), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(3)]
        for row in table.rows:
            for i, width in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = width

    def generate_html_table(self) -> str:
        """Generate HTML table for web display"""
        if not self.assessments:
            return "<p>No GRADE assessments available</p>"

        html = ['<table class="grade-table">']
        html.append("<thead><tr>")
        headers = [
            "Outcome",
            "Studies",
            "Design",
            "Risk of Bias",
            "Inconsistency",
            "Indirectness",
            "Imprecision",
            "Publication Bias",
            "Starting Certainty",
            "Final Certainty",
            "Downgrade Reasons",
        ]
        for h in headers:
            html.append(f"<th>{h}</th>")
        html.append("</tr></thead><tbody>")

        for a in self.assessments:
            certainty_class = a.final_certainty.value.lower().replace(" ", "-")
            html.append("<tr>")
            html.append(f"<td>{a.outcome}</td>")
            html.append(f"<td>{a.studies}</td>")
            html.append(f"<td>{a.design}</td>")
            html.append(f"<td>{a.risk_of_bias}</td>")
            html.append(f"<td>{a.inconsistency}</td>")
            html.append(f"<td>{a.indirectness}</td>")
            html.append(f"<td>{a.imprecision}</td>")
            html.append(f"<td>{a.publication_bias}</td>")
            html.append(f"<td>{a.starting_certainty.value}</td>")
            html.append(f'<td class="certainty-{certainty_class}">{a.final_certainty.value}</td>')
            reasons = "; ".join(a.downgrade_summary) if a.downgrade_summary else "None"
            html.append(f"<td>{reasons}</td>")
            html.append("</tr>")

        html.append("</tbody></table>")
        return "\n".join(html)


def create_grade_table(project_data: Dict[str, Any]) -> GRADETableGenerator:
    """Factory function to create GRADE table from project data"""
    generator = GRADETableGenerator(project_data)
    generator.auto_populate_from_meta()
    return generator


# Example usage
if __name__ == "__main__":
    # Demo
    sample_data = {
        "extraction": {
            "studies": [
                {"design": "RCT", "primary_outcome": "Mortality"},
                {"design": "RCT", "primary_outcome": "Mortality"},
            ]
        },
        "rob": {
            "assessments": [
                {"data": {"overall": "Low"}},
                {"data": {"overall": "Some concerns"}},
            ]
        },
        "meta": {
            "results": {
                "studies": [
                    {"design": "RCT", "int_n": 50, "ctrl_n": 50},
                    {"design": "RCT", "int_n": 60, "ctrl_n": 60},
                ],
                "heterogeneity": {"i2": 30},
                "pooled": {"effect": 0.75, "ci_lower": 0.55, "ci_upper": 1.02},
                "publication_bias": {"egger": {"significant": False}},
            }
        },
    }

    generator = create_grade_table(sample_data)
    print(json.dumps(generator.to_dict(), indent=2))
