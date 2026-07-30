"""
Export module - Word document, LaTeX, PRISMA reports
"""

import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pylatex import Document as LatexDocument, Section, Subsection, Command, Package, Figure, NoEscape, Tabular, MultiColumn
    from pylatex.utils import bold, italic
    LATEX_AVAILABLE = True
except ImportError:
    LATEX_AVAILABLE = False

from poolr.plotting.figures import create_forest_plot, create_funnel_plot, create_prisma_flow_diagram, save_figure


def export_to_word(project_data: Dict[str, Any], project_path: Path, output_path: Optional[Path] = None) -> Path:
    """Export full SRMA report to Word document"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Systematic Review and Meta-Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta = project_data.get("metadata", {})
    p = doc.add_paragraph()
    p.add_run(f"Generated: ").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
    p = doc.add_paragraph()
    p.add_run(f"Project Version: ").bold = True
    p.add_run(meta.get("version", "0.3.0"))
    
    doc.add_page_break()
    
    # 1. ABSTRACT
    doc.add_heading('Abstract', level=1)
    _add_abstract(doc, project_data)
    
    # 2. INTRODUCTION
    doc.add_heading('Introduction', level=1)
    _add_introduction(doc, project_data)
    
    # 3. METHODS
    doc.add_heading('Methods', level=1)
    _add_methods(doc, project_data)
    
    # 4. RESULTS
    doc.add_heading('Results', level=1)
    _add_results(doc, project_data)
    
    # 5. DISCUSSION
    doc.add_heading('Discussion', level=1)
    _add_discussion(doc, project_data)
    
    # 6. REFERENCES
    doc.add_heading('References', level=1)
    _add_references(doc, project_data)
    
    # Appendices
    doc.add_page_break()
    doc.add_heading('Appendices', level=1)
    _add_appendices(doc, project_data)
    
    # Save
    if output_path is None:
        output_path = project_path / "SRMA_Report.docx"
    doc.save(str(output_path))
    return output_path


def _add_abstract(doc, project_data):
    """Add structured abstract"""
    meta_results = project_data.get("meta", {}).get("results", {})
    
    abstract_text = (
        "Background: "
    )
    pico = project_data.get("pico", {})
    if pico.get("population"):
        abstract_text += f" {pico['population']}. "
    
    abstract_text += (
        "We conducted a systematic review and meta-analysis to evaluate "
    )
    if pico.get("intervention"):
        abstract_text += f"the effect of {pico['intervention'].lower()} "
    if pico.get("comparator"):
        abstract_text += f"compared to {pico['comparator'].lower()} "
    if pico.get("outcomes"):
        abstract_text += f"on {pico['outcomes'].lower()}. "
    
    doc.add_paragraph(abstract_text)
    
    if meta_results:
        p = doc.add_paragraph()
        p.add_run("Methods: ").bold = True
        p.add_run(f"We searched {', '.join(project_data.get('protocol', {}).get('databases', 'databases').split(', '))} "
                  f"from {project_data.get('protocol', {}).get('date_range', 'inception')}. ")
        
        p = doc.add_paragraph()
        p.add_run("Results: ").bold = True
        p.add_run(f"{len(project_data.get('screening', {}).get('title_abstract', []))} records were screened, "
                  f"{len(project_data.get('screening', {}).get('full_text', []))} assessed for eligibility, "
                  f"and {len(project_data.get('extraction', {}).get('studies', []))} studies included. ")
        
        pooled = meta_results.get("pooled", {})
        if pooled:
            measure = meta_results.get("measure", "OR")
            effect = pooled.get("effect", 0)
            ci_l = pooled.get("ci_lower", 0)
            ci_u = pooled.get("ci_upper", 0)
            p.add_run(f"Pooled {measure} = {effect:.2f} (95% CI: {ci_l:.2f}–{ci_u:.2f}). ")
        
        hetero = meta_results.get("heterogeneity", {})
        if hetero:
            p.add_run(f"I² = {hetero.get('i2', 0):.1f}%. ")
        
        p = doc.add_paragraph()
        p.add_run("Conclusions: ").bold = True
        if pooled:
            if pooled.get("ci_lower", 1) > 1 or pooled.get("ci_upper", 1) < 1:
                p.add_run("The intervention showed a statistically significant effect. ")
            else:
                p.add_run("No statistically significant difference was found. ")


def _add_introduction(doc, project_data):
    """Add introduction section"""
    pico = project_data.get("pico", {})
    
    doc.add_heading('Background', level=2)
    doc.add_paragraph(
        "Describe the rationale for the review in the context of what is already known. "
        "This should be customized based on your specific topic."
    )
    
    doc.add_heading('Objectives', level=2)
    objectives = (
        f"To evaluate the effectiveness of {pico.get('intervention', '[intervention]').lower()} "
        f"compared to {pico.get('comparator', '[comparator]').lower()} "
        f"for {pico.get('population', '[population]').lower()} "
        f"on outcomes including {pico.get('outcomes', '[outcomes]').lower()}."
    )
    doc.add_paragraph(objectives)
    
    # PICO table
    doc.add_heading('PICO Criteria', level=2)
    pico_table = doc.add_table(rows=5, cols=2)
    pico_table.style = 'Table Grid'
    pico_data = [
        ("Element", "Description"),
        ("Population", pico.get("population", "Not specified")),
        ("Intervention", pico.get("intervention", "Not specified")),
        ("Comparator", pico.get("comparator", "Not specified")),
        ("Outcomes", pico.get("outcomes", "Not specified")),
    ]
    for i, (elem, desc) in enumerate(pico_data):
        pico_table.rows[i].cells[0].text = elem
        pico_table.rows[i].cells[1].text = desc
        if i == 0:
            for cell in pico_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _add_methods(doc, project_data):
    """Add methods section"""
    protocol = project_data.get("protocol", {})
    
    doc.add_heading('Protocol and Registration', level=2)
    reg = protocol.get("registration", "Not registered")
    doc.add_paragraph(f"Protocol registered at: {reg}")
    
    doc.add_heading('Eligibility Criteria', level=2)
    doc.add_paragraph("Studies were included if they met the following criteria:")
    
    # Inclusion criteria
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Study design: ").bold = True
    p.add_run(protocol.get("study_designs", "RCTs, cohort studies"))
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Population: ").bold = True
    p.add_run(pico.get("population", "As per PICO"))
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Intervention: ").bold = True
    p.add_run(pico.get("intervention", "As per PICO"))
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Comparator: ").bold = True
    p.add_run(pico.get("comparator", "As per PICO"))
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Outcomes: ").bold = True
    p.add_run(pico.get("outcomes", "As per PICO"))
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Language: ").bold = True
    p.add_run(protocol.get("languages", "English"))
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Date range: ").bold = True
    p.add_run(protocol.get("date_range", "All years"))
    
    doc.add_heading('Information Sources', level=2)
    databases = protocol.get("databases", "PubMed, Embase, Cochrane CENTRAL, Scopus")
    doc.add_paragraph(f"Databases searched: {databases}")
    doc.add_paragraph("Search dates: [Date of search]")
    
    doc.add_heading('Search Strategy', level=2)
    doc.add_paragraph("Full search strategies for each database are provided in Appendix A.")
    
    doc.add_heading('Selection Process', level=2)
    doc.add_paragraph(
        "Two reviewers independently screened titles/abstracts and full texts. "
        "Disagreements were resolved by consensus or a third reviewer."
    )
    
    doc.add_heading('Data Collection Process', level=2)
    doc.add_paragraph(
        "Two reviewers independently extracted data using a standardized form. "
        "Discrepancies were resolved by discussion."
    )
    
    doc.add_heading('Data Items', level=2)
    doc.add_paragraph(
        "The following data were extracted: study characteristics, participant demographics, "
        "intervention details, comparator details, outcome measures, results, and risk of bias assessments."
    )
    
    doc.add_heading('Risk of Bias Assessment', level=2)
    rob_tools = []
    assessments = project_data.get("rob", {}).get("assessments", [])
    for a in assessments:
        if a.get("tool") not in rob_tools:
            rob_tools.append(a.get("tool"))
    if rob_tools:
        doc.add_paragraph(f"Risk of bias was assessed using: {', '.join(rob_tools)}.")
    else:
        doc.add_paragraph("Risk of bias was assessed using RoB 2 for RCTs, NOS for cohort studies, and PROBAST for diagnostic/prognostic studies.")
    
    doc.add_heading('Effect Measures', level=2)
    meta_results = project_data.get("meta", {}).get("results", {})
    measure = meta_results.get("measure", "OR")
    doc.add_paragraph(f"Primary effect measure: {measure} with 95% confidence intervals.")
    
    doc.add_heading('Synthesis Methods', level=2)
    model = meta_results.get("model", "Random-effects")
    method = meta_results.get("method", "DerSimonian-Laird")
    doc.add_paragraph(f"Meta-analysis was performed using a {model.lower()} model with the {method} estimator.")
    
    if meta_results.get("subgroups"):
        doc.add_paragraph(f"Subgroup analyses were performed by {project_data.get('meta', {}).get('subgroup', 'study design')}.")
    
    if meta_results.get("publication_bias"):
        doc.add_paragraph("Publication bias was assessed using Egger's test and funnel plot asymmetry.")
    
    if meta_results.get("meta_regression"):
        doc.add_paragraph("Meta-regression was performed to explore sources of heterogeneity.")
    
    doc.add_heading('Certainty Assessment', level=2)
    doc.add_paragraph("The certainty of evidence was assessed using the GRADE approach (see Appendix B).")


def _add_results(doc, project_data):
    """Add results section"""
    # Study selection
    doc.add_heading('Study Selection', level=2)
    ta = len(project_data.get("screening", {}).get("title_abstract", []))
    ft = len(project_data.get("screening", {}).get("full_text", []))
    included = len(project_data.get("extraction", {}).get("studies", []))
    
    doc.add_paragraph(
        f"The search yielded {ta} records. After removing duplicates, {ta} records were screened "
        f"by title/abstract. Of these, {ft} full-text reports were assessed for eligibility, "
        f"and {included} studies met inclusion criteria and were included in the review."
    )
    
    doc.add_paragraph("See PRISMA flow diagram (Figure 1) for details.")
    
    # Add PRISMA flow diagram if available
    if project_data.get("prisma"):
        try:
            from poolr.plotting.figures import create_prisma_flow_diagram, save_figure
            flow = project_data.get("prisma", {}).get("flow", {})
            fig = create_prisma_flow_diagram(flow)
            img_path = project_data.get("project_path", Path(".")) / "prisma_flow.png"
            save_figure(fig, img_path)
            doc.add_picture(str(img_path), width=Inches(6.5))
            doc.add_paragraph("Figure 1. PRISMA 2020 flow diagram.", style='Caption')
        except Exception:
            doc.add_paragraph("[PRISMA flow diagram could not be generated]")
    
    # Study characteristics
    doc.add_heading('Study Characteristics', level=2)
    studies = project_data.get("extraction", {}).get("studies", [])
    if studies:
        _add_study_characteristics_table(doc, studies)
    else:
        doc.add_paragraph("No study characteristics available.")
    
    # Risk of bias
    doc.add_heading('Risk of Bias Assessment', level=2)
    assessments = project_data.get("rob", {}).get("assessments", [])
    if assessments:
        _add_rob_summary(doc, assessments)
    else:
        doc.add_paragraph("No risk of bias assessments available.")
    
    # Results of syntheses
    doc.add_heading('Results of Syntheses', level=2)
    meta_results = project_data.get("meta", {}).get("results", {})
    if meta_results:
        _add_meta_results(doc, meta_results)
        
        # Forest plot
        try:
            fig = create_forest_plot(meta_results)
            img_path = project_data.get("project_path", Path(".")) / "forest_plot.png"
            save_figure(fig, img_path)
            doc.add_picture(str(img_path), width=Inches(6.5))
            doc.add_paragraph("Figure 2. Forest plot.", style='Caption')
        except Exception:
            doc.add_paragraph("[Forest plot could not be generated]")
        
        # Funnel plot
        if len(meta_results.get("studies", [])) >= 3:
            try:
                fig = create_funnel_plot(meta_results)
                img_path = project_data.get("project_path", Path(".")) / "funnel_plot.png"
                save_figure(fig, img_path)
                doc.add_picture(str(img_path), width=Inches(5))
                doc.add_paragraph("Figure 3. Funnel plot.", style='Caption')
            except Exception:
                doc.add_paragraph("[Funnel plot could not be generated]")
    else:
        doc.add_paragraph("No meta-analysis results available.")
    
    # Subgroup analysis
    if meta_results.get("subgroups"):
        doc.add_heading('Subgroup Analysis', level=2)
        for sg in meta_results["subgroups"]:
            p = doc.add_paragraph()
            p.add_run(f"{sg['name']}: ").bold = True
            p.add_run(f"{sg['measure']} = {sg['effect']:.2f} (95% CI: {sg['ci_lower']:.2f}–{sg['ci_upper']:.2f}), k={sg['k']}")
    
    # Publication bias
    if meta_results.get("publication_bias"):
        doc.add_heading('Publication Bias', level=2)
        for test, val in meta_results["publication_bias"].items():
            if isinstance(val, dict):
                p = doc.add_paragraph()
                p.add_run(f"{test.capitalize()}: ").bold = True
                p.add_run(f"p = {val.get('p_value', 'N/A'):.4f}")


def _add_study_characteristics_table(doc, studies):
    """Add table of study characteristics"""
    if not studies:
        return
    
    # Key columns to show
    cols = ["study_id", "design", "country", "year", "sample_size", "int_n", "ctrl_n", "int_events", "ctrl_events"]
    headers = ["Study", "Design", "Country", "Year", "Total N", "Int N", "Ctrl N", "Int Events", "Ctrl Events"]
    
    table = doc.add_table(rows=len(studies) + 1, cols=len(cols))
    table.style = 'Table Grid'
    
    # Headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
    
    # Data
    for row_idx, study in enumerate(studies):
        for col_idx, key in enumerate(cols):
            value = study.get(key, "")
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(1.5)
        row.cells[2].width = Cm(1.5)


def _add_rob_summary(doc, assessments):
    """Add risk of bias summary"""
    doc.add_paragraph(f"Risk of bias was assessed for {len(assessments)} studies using multiple tools.")
    
    # Group by tool
    tools = {}
    for a in assessments:
        tool = a.get("tool", "Unknown")
        if tool not in tools:
            tools[tool] = []
        tools[tool].append(a)
    
    for tool, tool_assessments in tools.items():
        doc.add_heading(f'{tool} Assessment', level=3)
        
        if tool == "RoB 2":
            domains = ["randomization", "deviations", "missing_data", "measurement", "selection", "overall"]
            domain_labels = ["Randomization", "Deviations", "Missing Data", "Measurement", "Selection", "Overall"]
        elif tool == "NOS":
            domains = ["representativeness", "selection", "ascertainment", "outcome", "comparability1", "comparability2", "assessment", "followup", "adequacy"]
            domain_labels = ["Representativeness", "Selection", "Ascertainment", "Outcome", "Comp 1", "Comp 2", "Assessment", "Follow-up", "Adequacy"]
        elif tool == "PROBAST":
            domains = ["participants", "predictors", "outcome", "analysis", "app_participants", "app_predictors", "app_outcome"]
            domain_labels = ["Participants", "Predictors", "Outcome", "Analysis", "App Part.", "App Pred.", "App Out."]
        else:
            domains = []
            domain_labels = []
        
        if domains:
            table = doc.add_table(rows=len(tool_assessments) + 1, cols=len(domains) + 1)
            table.style = 'Table Grid'
            
            # Headers
            table.rows[0].cells[0].text = "Study"
            for i, label in enumerate(domain_labels):
                table.rows[0].cells[i + 1].text = label
                for p in table.rows[0].cells[i + 1].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(7)
            
            for row_idx, assessment in enumerate(tool_assessments):
                data = assessment.get("data", {})
                study_id = assessment.get("study_id", f"Study {row_idx + 1}")
                table.rows[row_idx + 1].cells[0].text = study_id
                for i, domain in enumerate(domains):
                    table.rows[row_idx + 1].cells[i + 1].text = data.get(domain, "")
                    for p in table.rows[row_idx + 1].cells[i + 1].paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(7)


def _add_meta_results(doc, results):
    """Add meta-analysis results"""
    pooled = results.get("pooled", {})
    hetero = results.get("heterogeneity", {})
    measure = results.get("measure", "OR")
    model = results.get("model", "Random-effects")
    
    if pooled:
        p = doc.add_paragraph()
        p.add_run(f"Pooled {measure} ({model}): ").bold = True
        effect = pooled.get("effect", 0)
        ci_l = pooled.get("ci_lower", 0)
        ci_u = pooled.get("ci_upper", 0)
        p_val = pooled.get("p", 1)
        p.add_run(f"{effect:.2f} (95% CI: {ci_l:.2f}–{ci_u:.2f}), P = {p_val:.4f}")
        
        if hetero:
            p = doc.add_paragraph()
            p.add_run("Heterogeneity: ").bold = True
            p.add_run(f"Q = {hetero.get('q', 0):.2f} (df = {hetero.get('df', 0)}), P = {hetero.get('q_p', 1):.4f}; "
                      f"I² = {hetero.get('i2', 0):.1f}%; τ² = {hetero.get('tau2', 0):.4f}")
            
            # Interpretation
            i2 = hetero.get('i2', 0)
            if i2 < 25:
                p.add_run(" (Low heterogeneity)")
            elif i2 < 50:
                p.add_run(" (Moderate heterogeneity)")
            elif i2 < 75:
                p.add_run(" (Substantial heterogeneity)")
            else:
                p.add_run(" (Considerable heterogeneity)")
        
        # Significance
        p = doc.add_paragraph()
        if pooled.get("ci_lower", 1) > 1 or pooled.get("ci_upper", 1) < 1:
            p.add_run("The result is statistically significant (95% CI excludes null).")
        else:
            p.add_run("The result is not statistically significant (95% CI includes null).")
    
    # Individual study table
    if results.get("studies"):
        doc.add_heading('Individual Study Results', level=3)
        table = doc.add_table(rows=len(results["studies"]) + 1, cols=5)
        table.style = 'Table Grid'
        
        headers = ["Study", f"{measure}", "95% CI", "Weight (%)", "Subgroup"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            for p in table.rows[0].cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True
        
        for row_idx, s in enumerate(results["studies"]):
            table.rows[row_idx + 1].cells[0].text = s.get("study", "")
            table.rows[row_idx + 1].cells[1].text = f"{s.get('effect', 0):.2f}"
            table.rows[row_idx + 1].cells[2].text = f"{s.get('ci_lower', 0):.2f}–{s.get('ci_upper', 0):.2f}"
            table.rows[row_idx + 1].cells[3].text = f"{s.get('weight', 0):.1f}%"
            table.rows[row_idx + 1].cells[4].text = s.get("subgroup", "")


def _add_discussion(doc, project_data):
    """Add discussion section"""
    doc.add_heading('Summary of Evidence', level=2)
    doc.add_paragraph(
        "Summarize the main findings including the strength of evidence for each main outcome. "
        "Consider the GRADE assessment."
    )
    
    doc.add_heading('Limitations', level=2)
    doc.add_paragraph("Discuss limitations at study and outcome level (e.g., risk of bias), and at review level (e.g., incomplete retrieval of identified research, reporting bias).")
    
    doc.add_heading('Conclusions', level=2)
    doc.add_paragraph("Provide a general interpretation of the results in the context of other evidence, and implications for future research.")


def _add_references(doc, project_data):
    """Add references section"""
    doc.add_paragraph("References would be automatically populated from study citations in a full implementation.")
    studies = project_data.get("extraction", {}).get("studies", [])
    for i, s in enumerate(studies):
        ref = f"{i+1}. {s.get('authors', 'Unknown')}. {s.get('title', 'Untitled')}. "
        ref += f"{s.get('journal', 'Unknown')}. {s.get('year', 'Unknown')}"
        if s.get('doi'):
            ref += f". doi:{s['doi']}"
        doc.add_paragraph(ref, style='List Number')


def _add_appendices(doc, project_data):
    """Add appendices"""
    # Appendix A: Search strategies
    doc.add_heading('Appendix A: Search Strategies', level=2)
    strategies = project_data.get("search_strategies", {})
    for db, strat in strategies.items():
        doc.add_heading(db.upper(), level=3)
        doc.add_paragraph(strat)
    
    # Appendix B: GRADE table
    doc.add_heading('Appendix B: GRADE Evidence Profile', level=2)
    _add_grade_table(doc, project_data)
    
    # Appendix C: Excluded studies
    doc.add_heading('Appendix C: Excluded Studies', level=2)
    _add_excluded_studies(doc, project_data)


def _add_grade_table(doc, project_data):
    """Add GRADE evidence table"""
    meta_results = project_data.get("meta", {}).get("results", {})
    studies = project_data.get("extraction", {}).get("studies", [])
    rob_assessments = project_data.get("rob", {}).get("assessments", [])
    
    # Determine GRADE ratings (simplified)
    n_studies = len(studies)
    measure = meta_results.get("measure", "OR")
    hetero = meta_results.get("heterogeneity", {})
    i2 = hetero.get("i2", 0)
    pooled = meta_results.get("pooled", {})
    
    # Starting rating
    designs = set(s.get("design", "") for s in studies)
    if "RCT" in designs and len(designs) == 1:
        starting = "High"
    else:
        starting = "Low"
    
    # Downgrades
    downgrades = []
    
    # Risk of bias
    if rob_assessments:
        high_rob = sum(1 for a in rob_assessments if a.get("data", {}).get("overall") == "High")
        if high_rob / len(rob_assessments) > 0.25:
            downgrades.append("Risk of bias: Serious concerns")
    
    # Inconsistency
    if i2 > 50:
        downgrades.append("Inconsistency: Substantial heterogeneity (I² > 50%)")
    elif i2 > 25:
        downgrades.append("Inconsistency: Moderate heterogeneity (I² > 25%)")
    
    # Indirectness
    downgrades.append("Indirectness: No serious concerns (direct evidence)")
    
    # Imprecision
    if pooled:
        ci_l = pooled.get("ci_lower", 0)
        ci_u = pooled.get("ci_upper", 0)
        if ci_l < 0.5 and ci_u > 2:  # Wide CI crossing threshold
            downgrades.append("Imprecision: Wide confidence interval")
    
    # Publication bias
    pub_bias = meta_results.get("publication_bias", {})
    if pub_bias.get("egger", {}).get("significant", False):
        downgrades.append("Publication bias: Egger's test significant")
    
    # Final rating
    final = starting
    if len(downgrades) >= 2:
        final = "Very Low" if starting == "Low" else "Low"
    elif len(downgrades) == 1:
        final = "Low" if starting == "High" else "Very Low"
    
    # Create table
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    
    headers = ["Domain", "Judgment", "Downgrade", "Comments"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    
    data = [
        ("Study design", starting, "", "Based on included study designs"),
        ("Risk of bias", "Serious" if any("Risk of bias" in d for d in downgrades) else "Not serious", 
         "↓" if any("Risk of bias" in d for d in downgrades) else "", 
         next((d for d in downgrades if "Risk of bias" in d), "No serious concerns")),
        ("Inconsistency", "Serious" if any("Inconsistency" in d for d in downgrades) else "Not serious",
         "↓" if any("Inconsistency" in d for d in downgrades) else "",
         next((d for d in downgrades if "Inconsistency" in d), "No serious concerns")),
        ("Indirectness", "Not serious", "", "Direct evidence"),
        ("Imprecision", "Serious" if any("Imprecision" in d for d in downgrades) else "Not serious",
         "↓" if any("Imprecision" in d for d in downgrades) else "",
         next((d for d in downgrades if "Imprecision" in d), "No serious concerns")),
        ("Publication bias", "Serious" if any("Publication bias" in d for d in downgrades) else "Not serious",
         "↓" if any("Publication bias" in d for d in downgrades) else "",
         next((d for d in downgrades if "Publication bias" in d), "No serious concerns")),
        ("Final", final, "", f"{len(downgrades)} downgrade(s) from {starting}"),
    ]
    
    for row_idx, (domain, judgment, downgrade, comment) in enumerate(data):
        table.rows[row_idx + 1].cells[0].text = domain
        table.rows[row_idx + 1].cells[1].text = judgment
        table.rows[row_idx + 1].cells[2].text = downgrade
        table.rows[row_idx + 1].cells[3].text = comment
    
    doc.add_paragraph(f"\nFinal GRADE rating: {final}")


def _add_excluded_studies(doc, project_data):
    """Add table of excluded studies"""
    ft_records = project_data.get("screening", {}).get("full_text", [])
    excluded = [r for r in ft_records if r.get("decision_reviewer1") is False or r.get("decision_reviewer2") is False]
    
    if not excluded:
        doc.add_paragraph("No full-text exclusions recorded.")
        return
    
    table = doc.add_table(rows=len(excluded) + 1, cols=4)
    table.style = 'Table Grid'
    
    headers = ["Study", "Reviewer", "Decision", "Reason"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    
    for row_idx, rec in enumerate(excluded):
        r1 = rec.get("decision_reviewer1")
        r2 = rec.get("decision_reviewer2")
        
        reviewer = []
        if r1 is False:
            reviewer.append("R1")
        if r2 is False:
            reviewer.append("R2")
        
        reason = rec.get("reason_reviewer1") or rec.get("reason_reviewer2") or ""
        
        table.rows[row_idx + 1].cells[0].text = rec.get("title", "Unknown")[:50]
        table.rows[row_idx + 1].cells[1].text = ", ".join(reviewer)
        table.rows[row_idx + 1].cells[2].text = "Excluded"
        table.rows[row_idx + 1].cells[3].text = reason[:100]


def export_to_latex(project_data: Dict[str, Any], project_path: Path, output_path: Optional[Path] = None) -> Path:
    """Export to LaTeX (requires manual compilation)"""
    if not LATEX_AVAILABLE:
        raise ImportError("pylatex not installed. Install with: pip install pylatex")
    
    doc = LatexDocument(documentclass='article', document_options=['12pt', 'a4paper'])
    doc.packages.append(Package('geometry', options=['margin=1in']))
    doc.packages.append(Package('booktabs'))
    doc.packages.append(Package('graphicx'))
    doc.packages.append(Package('float'))
    doc.packages.append(Package('hyperref'))
    doc.packages.append(Package('longtable'))
    
    # Title
    doc.preamble.append(Command('title', 'Systematic Review and Meta-Analysis Report'))
    doc.preamble.append(Command('author', 'Generated by poolr'))
    doc.preamble.append(Command('date', datetime.now().strftime('%B %d, %Y')))
    doc.append(NoEscape(r'\maketitle'))
    
    # Abstract
    with doc.create(Section('Abstract')):
        doc.append(_generate_abstract_text(project_data))
    
    # Introduction
    with doc.create(Section('Introduction')):
        with doc.create(Subsection('Background')):
            doc.append("Background rationale...")
        with doc.create(Subsection('Objectives')):
            pico = project_data.get("pico", {})
            doc.append(f"To evaluate {pico.get('intervention', '[intervention]')} vs {pico.get('comparator', '[comparator]')} "
                      f"for {pico.get('population', '[population]')} on {pico.get('outcomes', '[outcomes]')}. ")
    
    # Methods
    with doc.create(Section('Methods')):
        _add_latex_methods(doc, project_data)
    
    # Results
    with doc.create(Section('Results')):
        _add_latex_results(doc, project_data)
    
    # Discussion
    with doc.create(Section('Discussion')):
        doc.append("Discussion...")
    
    # Save
    if output_path is None:
        output_path = project_path / "SRMA_Report.tex"
    doc.generate_tex(str(output_path.with_suffix('')))
    return output_path


def _add_latex_methods(doc, project_data):
    """Add methods to LaTeX doc"""
    with doc.create(Subsection('Protocol and Registration')):
        doc.append(f"Protocol registered: {project_data.get('protocol', {}).get('registration', 'Not registered')}")
    
    with doc.create(Subsection('Eligibility Criteria')):
        doc.append("Studies were included if they met PICO criteria...")
    
    with doc.create(Subsection('Information Sources')):
        databases = project_data.get('protocol', {}).get('databases', 'PubMed, Embase, Cochrane CENTRAL, Scopus')
        doc.append(f"Databases: {databases}")
    
    with doc.create(Subsection('Search Strategy')):
        doc.append("Full search strategies in Appendix A.")
    
    with doc.create(Subsection('Selection Process')):
        doc.append("Two reviewers independently screened...")
    
    with doc.create(Subsection('Risk of Bias Assessment')):
        doc.append("RoB 2, NOS, PROBAST as appropriate.")
    
    with doc.create(Subsection('Effect Measures')):
        measure = project_data.get('meta', {}).get('results', {}).get('measure', 'OR')
        doc.append(f"Primary measure: {measure}")
    
    with doc.create(Subsection('Synthesis Methods')):
        model = project_data.get('meta', {}).get('results', {}).get('model', 'Random-effects')
        doc.append(f"Meta-analysis: {model} model.")


def _add_latex_results(doc, project_data):
    """Add results to LaTeX doc"""
    ta = len(project_data.get("screening", {}).get("title_abstract", []))
    ft = len(project_data.get("screening", {}).get("full_text", []))
    included = len(project_data.get("extraction", {}).get("studies", []))
    
    with doc.create(Subsection('Study Selection')):
        doc.append(f"Search yielded {ta} records. {ft} full-text assessed. {included} included.")
    
    meta_results = project_data.get("meta", {}).get("results", {})
    if meta_results:
        with doc.create(Subsection('Synthesis Results')):
            pooled = meta_results.get("pooled", {})
            measure = meta_results.get("measure", "OR")
            if pooled:
                effect = pooled.get("effect", 0)
                ci_l = pooled.get("ci_lower", 0)
                ci_u = pooled.get("ci_upper", 0)
                doc.append(f"Pooled {measure} = {effect:.2f} (95% CI: {ci_l:.2f}–{ci_u:.2f})")
            
            hetero = meta_results.get("heterogeneity", {})
            if hetero:
                doc.append(f"I² = {hetero.get('i2', 0):.1f}%, τ² = {hetero.get('tau2', 0):.4f}")


def _generate_abstract_text(project_data):
    """Generate abstract text"""
    pico = project_data.get("pico", {})
    meta_results = project_data.get("meta", {}).get("results", {})
    
    text = f"Background: {pico.get('population', 'Population')}... "
    text += f"We evaluated {pico.get('intervention', 'intervention')} vs {pico.get('comparator', 'comparator')} "
    text += f"on {pico.get('outcomes', 'outcomes')}. "
    text += "Methods: Systematic search of databases. "
    
    if meta_results:
        pooled = meta_results.get("pooled", {})
        if pooled:
            measure = meta_results.get("measure", "OR")
            effect = pooled.get("effect", 0)
            ci_l = pooled.get("ci_lower", 0)
            ci_u = pooled.get("ci_upper", 0)
            text += f"Results: Pooled {measure} = {effect:.2f} (95% CI: {ci_l:.2f}–{ci_u:.2f}). "
    
    text += "Conclusions: Further research needed."
    return text


# CLI entry points
def cli_main():
    """Headless CLI for CI/CD"""
    import argparse
    parser = argparse.ArgumentParser(description="poolr CLI - headless SRMA processing")
    parser.add_argument("project", help="Path to project folder")
    parser.add_argument("--run-meta", action="store_true", help="Run meta-analysis")
    parser.add_argument("--export", choices=["word", "latex", "json", "all"], default="json", help="Export format")
    parser.add_argument("--output", help="Output path")
    args = parser.parse_args()
    
    project_path = Path(args.project)
    poolr_json = project_path / "poolr.json"
    
    if not poolr_json.exists():
        print(f"Error: No poolr.json found in {project_path}")
        return 1
    
    with open(poolr_json, "r") as f:
        project_data = json.load(f)
    
    project_data["project_path"] = project_path
    
    if args.run_meta:
        from poolr.meta.analysis import MetaAnalysis
        studies = project_data.get("extraction", {}).get("studies", [])
        if studies:
            print("Running meta-analysis...")
            # Auto-detect measure from data
            has_binary = any(s.get("int_events") and s.get("ctrl_events") for s in studies)
            has_continuous = any(s.get("int_mean") and s.get("ctrl_mean") for s in studies)
            
            if has_binary:
                meta = MetaAnalysis(model="random", measure="OR")
                results = meta.run([s for s in studies if s.get("int_events")])
                project_data["meta"]["results"] = results
                print(f"Pooled OR = {results['pooled']['effect']:.3f} "
                      f"(95% CI: {results['pooled']['ci_lower']:.3f}-{results['pooled']['ci_upper']:.3f})")
                print(f"I² = {results['heterogeneity']['i2']:.1f}%")
    
    if args.export in ["word", "all"]:
        if DOCX_AVAILABLE:
            out = export_to_word(project_data, project_path, Path(args.output) if args.output else None)
            print(f"Word report saved to: {out}")
        else:
            print("python-docx not installed, skipping Word export")
    
    if args.export in ["latex", "all"]:
        if LATEX_AVAILABLE:
            out = export_to_latex(project_data, project_path, Path(args.output) if args.output else None)
            print(f"LaTeX report saved to: {out}")
        else:
            print("pylatex not installed, skipping LaTeX export")
    
    if args.export in ["json", "all"]:
        out = project_path / "poolr_report.json"
        with open(out, "w") as f:
            json.dump(project_data, f, indent=2)
        print(f"JSON report saved to: {out}")
    
    # Save updated project
    with open(poolr_json, "w") as f:
        json.dump(project_data, f, indent=2)
    
    return 0


if __name__ == "__main__":
    cli_main()